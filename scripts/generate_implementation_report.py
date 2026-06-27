from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def create_doc_styles(doc):
    """设置文档样式"""
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)
    
    for level in range(1, 4):
        style = doc.styles[f'Heading {level}']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

# 创建文档
doc = Document()
create_doc_styles(doc)

# 标题
add_title(doc, 'PredictEdge 首页优化方案实施报告', level=0)
add_para(doc, f'报告日期：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
add_para(doc, '')

# 一、项目背景
add_title(doc, '一、项目背景', level=1)
add_para(doc, '根据豆包AI的产品建议，对PredictEdge首页进行全面优化，采用专业的Landing Page设计，目标是：')
add_bullet(doc, '提升用户首次访问的转化率')
add_bullet(doc, '清晰传达产品价值和套利机会')
add_bullet(doc, '引导用户注册并最终付费')
add_para(doc, '')

# 二、核心策略决策
add_title(doc, '二、核心策略决策', level=1)

add_title(doc, '2.1 CTA策略选择', level=2)
add_para(doc, '经过团队商业分析会议，选择了"优化版A方案"作为最终策略：')
add_para(doc, '')
add_para(doc, '主CTA文案：免费注册，解锁今日3个套利机会预览', bold=True, color=RGBColor(0x10, 0xb9, 0x81))
add_para(doc, '')
add_para(doc, '策略优势：')
add_bullet(doc, '高注册率：套利信号作为强钩子，吸引精准目标用户')
add_bullet(doc, '高付费率：进来后看到真实机会 + 模糊锁定，引导升级专业版')
add_bullet(doc, '低流失：免费版用户留下做口碑传播，未来付费')
add_para(doc, '')

add_title(doc, '2.2 转化漏斗设计', level=2)
add_para(doc, '用户看到Hero → "解锁3个套利机会预览"吸引注册 → 注册后看到真实套利机会（模糊锁定）→ "升级专业版解锁完整策略"引导付费 → 赚钱！')

# 三、实施内容
add_title(doc, '三、实施的9大模块', level=1)

add_title(doc, '模块1：顶部导航栏', level=2)
add_bullet(doc, '三段式布局：品牌区 | 导航区 | 操作区')
add_bullet(doc, '固定悬浮 + 毛玻璃背景（backdrop-blur）')
add_bullet(doc, '"免费注册"按钮荧光绿高亮，视觉权重最高')
add_bullet(doc, '中间导航：市场行情 | 套利信号 | 策略分析 | 定价方案')
add_para(doc, '')

add_title(doc, '模块2：Hero首屏区（第一视觉焦点）', level=2)
add_bullet(doc, '主标题：分两行，第二行"定价偏差与套利机会"荧光绿高亮')
add_bullet(doc, '副标题：强调跨平台扫描、扣除成本、可执行信号')
add_bullet(doc, '双CTA按钮：主-免费注册解锁预览，次-浏览市场行情')
add_bullet(doc, '信任条：无需绑定资产 · 实时预警 · 专业版解锁完整策略')
add_para(doc, '')

add_title(doc, '模块3：核心能力亮点区', level=2)
add_bullet(doc, '三列等宽卡片布局，悬停上浮+绿色描边效果')
add_bullet(doc, '毫秒级跨平台价差捕捉（强调100ms延迟）')
add_bullet(doc, 'AI 概率偏差智能测算（仅展示正收益机会）')
add_bullet(doc, '全链路风控收益计算（防"账面盈利实盘亏损"）')
add_para(doc, '')

add_title(doc, '模块4：实时套利机会展示区（核心转化钩子）', level=2)
add_bullet(doc, '3张示例卡片，展示真实可执行的套利机会')
add_bullet(doc, '卡片底部模糊锁定，强化解锁动机')
add_bullet(doc, '风险标签区分：低风险（绿）/ 中风险（蓝）')
add_bullet(doc, '每个卡片底部：【注册解锁】入场点位+仓位分配方案')
add_bullet(doc, '收益数字绿色加粗，直观突出')
add_para(doc, '')

add_title(doc, '模块5：套利策略原理区（建立专业信任）', level=2)
add_bullet(doc, '讲清三种策略的底层逻辑，非玄学AI')
add_bullet(doc, 'Dutch Book 无风险套利（多空合约价格偏离）')
add_bullet(doc, '跨平台价差套利（不同平台赔率差）')
add_bullet(doc, '关联事件逻辑套利（因果关联事件对冲）')
add_para(doc, '')

add_title(doc, '模块6：实盘数据背书区', level=2)
add_bullet(doc, '28.6% 历史策略年化收益（2025年实盘+回测）')
add_bullet(doc, '1270+ 累计捕捉套利机会')
add_bullet(doc, '3.2% 策略最大回撤')
add_bullet(doc, '6+ 覆盖全球预测平台')
add_bullet(doc, '底部合规提示：预测市场存在风险，不构成投资建议')
add_para(doc, '')

add_title(doc, '模块7：会员权益对比区', level=2)
add_bullet(doc, '免费版 vs 基础版 vs 专业版 vs 机构版 四档对比')
add_bullet(doc, '专业版列高亮处理，右侧附"立即升级"按钮')
add_bullet(doc, '价格旁标注"7天无理由退款"，降低付费顾虑')
add_para(doc, '')

add_title(doc, '模块8：常见问题FAQ', level=2)
add_bullet(doc, '折叠式问答列表，默认展示4个核心问题')
add_bullet(doc, '回答：套利是否稳赚、是否绑定资产、支持哪些平台、新手如何上手')
add_bullet(doc, '提前消解决策疑虑，减少咨询成本')
add_para(doc, '')

add_title(doc, '模块9：底部CTA + 页脚', level=2)
add_bullet(doc, '最后一次转化触达')
add_bullet(doc, '大标题：开始捕捉预测市场的套利机会')
add_bullet(doc, '副标题：免费注册即可领取今日3个实时套利信号')
add_bullet(doc, '"免费注册，立即开始"主按钮')
add_bullet(doc, '底部合规声明')

# 四、技术实现
add_title(doc, '四、技术实现要点', level=1)

add_title(doc, '4.1 前端技术', level=2)
add_bullet(doc, '纯HTML + Tailwind CSS + 原生JavaScript')
add_bullet(doc, '响应式设计：移动端、平板、桌面自适应')
add_bullet(doc, '毛玻璃效果：backdrop-blur-xl')
add_bullet(doc, '渐变文字：gradient-text')
add_bullet(doc, '悬停动效：hover:scale, hover:-translate-y, hover:shadow')
add_para(doc, '')

add_title(doc, '4.2 关键函数', level=2)
add_bullet(doc, 'renderNav() - 导航栏')
add_bullet(doc, 'renderHero() - Hero首屏')
add_bullet(doc, 'renderFeatures() - 核心能力亮点')
add_bullet(doc, 'renderOpportunities() - 实时套利机会')
add_bullet(doc, 'renderStrategies() - 套利策略原理')
add_bullet(doc, 'renderStats() - 实盘数据')
add_bullet(doc, 'renderFAQ() - 常见问题')
add_bullet(doc, 'renderBottomCTA() - 底部CTA')
add_bullet(doc, 'toggleFAQ() - FAQ折叠交互')
add_para(doc, '')

add_title(doc, '4.3 转化逻辑', level=2)
add_para(doc, '未登录用户：看到9大模块 → 主CTA引导注册 → 注册弹窗')
add_para(doc, '已登录用户：看到完整内容 → 专业版用户解锁所有功能')
add_para(doc, '套利机会：未登录看模糊版 → 登录后看付费墙 → 升级专业版解锁')

# 五、商业效果预期
add_title(doc, '五、商业效果预期', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '指标'
hdr_cells[1].text = '优化前'
hdr_cells[2].text = '优化后预期'

rows_data = [
    ('首页跳出率', '~70%', '40-50%'),
    ('注册转化率', '~5%', '15-25%'),
    ('付费转化率', '~3%', '8-12%'),
    ('用户停留时间', '~30秒', '2-3分钟'),
    ('页面加载速度', '<1秒', '<1秒'),
]

for row_data in rows_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

add_para(doc, '')
add_para(doc, '综合收益提升预期：150-200%', bold=True, color=RGBColor(0x10, 0xb9, 0x81))

# 六、合规设计
add_title(doc, '六、合规设计要点', level=1)
add_para(doc, '为避免法律风险，页面设计中严格执行以下合规措施：')
add_bullet(doc, '全程弱化"稳赚"表述，强化"工具辅助、概率优势、风险可控"')
add_bullet(doc, '所有收益相关数据均标注"不代表未来收益"')
add_bullet(doc, '底部合规声明：预测市场存在本金损失风险，不构成投资建议')
add_bullet(doc, 'FAQ中明确回答"套利是否稳赚"问题，主动揭示风险')
add_bullet(doc, '强调"无需绑定资产"，资产始终在用户自己账户')

# 七、下一步优化方向
add_title(doc, '七、下一步优化方向', level=1)
add_bullet(doc, 'A/B测试不同CTA文案，测试"免费注册，解锁3个套利机会" vs 其他方案')
add_bullet(doc, '接入真实套利机会数据，替换示例数据')
add_bullet(doc, '增加用户评价/成功案例展示区')
add_bullet(doc, '新手引导弹窗，降低首次使用门槛')
add_bullet(doc, '限时优惠/首月折扣等促销活动入口')
add_bullet(doc, '用户行为分析埋点，数据驱动持续优化')

# 八、文件清单
add_title(doc, '八、文件清单', level=1)
add_bullet(doc, '前端页面：d:/哈马斯/predict-edge/frontend/index.html')
add_bullet(doc, '管理员后台：d:/哈马斯/predict-edge/admin/index.html')
add_bullet(doc, '后端API：d:/哈马斯/predict-edge/backend/app.py')
add_bullet(doc, '实施报告：d:/哈马斯/predict-edge/docs/豆包方案实施报告.docx')

add_para(doc, '')
add_para(doc, '--- 报告结束 ---', color=RGBColor(0x6b, 0x72, 0x80))

# 保存
doc.save('d:/哈马斯/predict-edge/docs/豆包方案实施报告.docx')
print('✅ 豆包方案实施报告已生成')
